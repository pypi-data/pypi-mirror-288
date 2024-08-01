def 合併文件(文件集, 合併文件檔=None, 文件以換頁符號分隔=True):
    '合併微軟辦公室文件(Word)。' 
    from docxcompose.composer import Composer
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from pathlib import Path
    文件集 = sorted(文件集)

    if not 合併文件檔:
        合併文件檔 = Path(文件集[0]).parent / '合併文件.docx'
    
    input_paths = 文件集
    # Create a new document based on the first input document
    merged_document = Document(input_paths[0])
    composer = Composer(merged_document)

    def 插入換頁符號(document):
        # Create a page break element
        page_break = OxmlElement('w:br')
        page_break.set(qn('w:type'), 'page')
        # Add the page break to the last paragraph
        document.paragraphs[-1]._p.append(page_break)

    for input_path in input_paths[1:]:
        if 文件以換頁符號分隔:
            插入換頁符號(merged_document)
        # Append each subsequent document
        composer.append(Document(input_path))

    # Save the merged document to the specified output path
    composer.save(str(合併文件檔))

def doc2docx(docs):
    '微軟辦公室文件為 doc 格式者，轉成 docx 格式。'
    import os
    import comtypes.client
    import pypandoc
    from collections.abc import Iterable 
    if isinstance(docs, str) and not isinstance(docs, Iterable):
        docs = [docs]
    for doc in docs:
        input_path = str(doc)
        output_path = str(doc.with_suffix('.docx'))
        # Ensure the input file is a .doc file
        if not input_path.lower().endswith('.doc'):
            raise ValueError("Input file must be a .doc file")

        # Ensure the output file is a .docx file
        if not output_path.lower().endswith('.docx'):
            raise ValueError("Output file must be a .docx file")

        # Create a Word application object
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        # Open the .doc file
        doc = word.Documents.Open(input_path)

        # Save the file as .docx
        doc.SaveAs(output_path, FileFormat=16)  # 16 corresponds to the wdFormatDocumentDefault format

        # Close the document and quit the application
        doc.Close()
        word.Quit()


from pathlib import Path
fs = list(Path(__file__).parent.glob('*.docx'))
合併文件(fs)
# print(fs)
# doc2docx(fs)
